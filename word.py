import docx

doc = docx.Document('/Users/user/w/python/app_local/operation_news/genkou.docx')

txt = []
str = ''

for i in range(len(doc.paragraphs)):
    list_item = doc.paragraphs[i].text
    if i == 0:
        str += list_item + '\n'
    elif i == 1:
        str += list_item + '\n'
    elif i == 2:
        str += list_item + '\n'
    elif i == 3:
        str += list_item + '\n'
    else:
        str += list_item + '<br>' + '\n'


'''
for par in doc.paragraphs:
    if par.text == doc.paragraphs[0].text:
        str += par.text + '\n'
    elif par.text == doc.paragraphs[2].text:
        str += par.text + '\n'
    else:
        par.add_run('<br>')
        str += par.text + '\n'
'''


print(str)

file = open('word.txt', 'w')
file.write(str)
file.close()










'''
file = open('InputText.txt', 'w')
file.write('asdfasd')
file.close()
'''
